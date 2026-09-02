"""Assemble the VTU Phase-1 final report as a .docx from the Markdown chapters.

The department submits Word documents, and the chapters are written in Markdown
so they can be diffed and reviewed. This script is the bridge. It applies the
VTU formatting conventions - Times New Roman 12 pt, 1.5 line spacing, justified
body, numbered headings, figures centred with numbered captions - so that the
formatting is a property of the build and not of whoever last edited the file.

    python docs/build_report.py                  # -> docs/report/EdgeGrid_Phase1_Report.docx
    python docs/build_report.py --open           # and open it
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
REPORT_DIR = HERE / "report"
FIG_DIR = HERE / "figures"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"

TITLE = "DePIN-Edge: A Decentralized Physical Infrastructure Network for Localized, Verifiable AI Inference"
SUBTITLE = "The Edge Grid"
TEAM = [
    ("HARSHIT RAJ", "1MV23IC021"),
    ("CHETAN RAGHUVANSHI", "1MV23IC013"),
    ("KESHAV NARAYAN", "1MV23IC023"),
    ("MAYUR AGARWAL", "1MV23IC028"),
]
GUIDE = "Dr. SAVITA CHOUDHARY"
GUIDE_TITLE = "Professor & Head"
DEPT = "Department of Computer Science and Engineering\n(IoT, Cyber-Security and Blockchain Technology)"
COLLEGE = "SIR M. VISVESVARAYA INSTITUTE OF TECHNOLOGY"
COLLEGE_TC = "Sir M. Visvesvaraya Institute of Technology"
COLLEGE_ADDR = ("Krishnadevaraya Nagar, International Airport Road,\n"
                "Hunasmaranahalli, Bengaluru - 562157")
UNIVERSITY = "VISVESVARAYA TECHNOLOGICAL UNIVERSITY, BELAGAVI"
UNIVERSITY_ADDR = "Jnana Sangama, Belagavi - 590 018"
YEAR = "2026 - 2027"
YEAR_SHORT = "2026-27"

# The department template names these three; they are constants here so a
# correction is a one-line edit rather than a search through prose.
PRINCIPAL = "Dr. M. N. Thippeswamy"
COORDINATOR = "Mr. Madhusudhana V"
COORDINATOR_TITLE = "Assistant Professor"
DEPT_SHORT = "Computer Science and Engineering (IoT, Cyber-Security and Blockchain Technology)"

ABSTRACT = (
    "The inference capacity that modern artificial intelligence depends upon is concentrated "
    "in a small number of hyperscale data centres, which imposes cost, latency, privacy and "
    "availability constraints on every application built above it, while a large installed "
    "base of consumer hardware sits idle. This project presents The Edge Grid, a decentralized "
    "physical infrastructure network in which an inference request is discovered, auctioned, "
    "executed, committed, verified and settled without a central directory, scheduler or "
    "payment rail. The system composes five mechanisms that the literature has studied "
    "separately: Kademlia distributed hash table discovery and a GossipSub task mempool over "
    "libp2p, a sealed-bid second-price auction with a warm-start incentive, a streaming "
    "inference runtime on commodity hardware, an LLM-as-a-Judge verification layer over "
    "Merkle-committed output blobs, and staked settlement on an Ethereum Virtual Machine "
    "chain with an escrow state machine and fraud-proof slashing. The contribution is the "
    "composition and its measurement rather than any individual mechanism. Evaluation on "
    "commodity CPU hardware records a mean warm time-to-first-token of 609.6 ms with every "
    "trial below one second, a cold-start penalty of 12.18 times that establishes the "
    "economic case for the warm-start incentive, auction clearing in tens of milliseconds "
    "across three to five peers, and settlement resolved correctly along all three of its "
    "paths on a deployed chain. The verification layer is measured to reject no honest "
    "output while detecting only 30 per cent of semantically inverted fraud, and to change "
    "its verdict on a quarter of claims that are merely reworded; the report argues that a "
    "judge drawn from the same model family as the provider inherits the errors it is "
    "deployed to police, and that the economic security of such a network is therefore "
    "bounded by the adversary's best strategy rather than the judge's average performance."
)

ABBREVIATIONS = [
    ("API", "Application Programming Interface"),
    ("CID", "Content Identifier (IPFS)"),
    ("CPU", "Central Processing Unit"),
    ("CSV", "Comma-Separated Values"),
    ("DA", "Data Availability"),
    ("DAO", "Decentralized Autonomous Organization"),
    ("DePIN", "Decentralized Physical Infrastructure Network"),
    ("DHT", "Distributed Hash Table"),
    ("ECDSA", "Elliptic Curve Digital Signature Algorithm"),
    ("EVM", "Ethereum Virtual Machine"),
    ("F1", "Harmonic mean of precision and recall"),
    ("FPR", "False Positive Rate"),
    ("GPU", "Graphics Processing Unit"),
    ("GRID", "The network's unit of account"),
    ("HTLC", "Hashed Timelock Contract"),
    ("IPFS", "InterPlanetary File System"),
    ("JSON", "JavaScript Object Notation"),
    ("L2", "Layer 2 (blockchain scaling layer)"),
    ("LLM", "Large Language Model"),
    ("LRU", "Least Recently Used"),
    ("NPU", "Neural Processing Unit"),
    ("P2P", "Peer-to-Peer"),
    ("RPC", "Remote Procedure Call"),
    ("SDK", "Software Development Kit"),
    ("SSE", "Server-Sent Events"),
    ("TTFT", "Time To First Token"),
    ("UDP", "User Datagram Protocol"),
    ("VRAM", "Video Random Access Memory"),
    ("VTU", "Visvesvaraya Technological University"),
    ("zkML", "Zero-Knowledge Machine Learning"),
]

# chapters in submission order
CHAPTER_FILES = [
    "ch1_3_introduction.md",
    "ch4_6_literature.md",
    "ch7_proposed_system.md",
    "ch8_results.md",          # written once the experiments have run
    "ch9_conclusion.md",
]


# ----------------------------------------------------------------- helpers

def _style(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = BODY_FONT
    n.font.size = Pt(12)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = n.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, bold, before, after in (
        ("Heading 1", 16, True, 18, 12),
        ("Heading 2", 14, True, 14, 8),
        ("Heading 3", 12.5, True, 12, 6),
        ("Heading 4", 12, True, 10, 6),
    ):
        s = doc.styles[name]
        s.font.name = BODY_FONT
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.keep_with_next = True


def _centred(doc, text, size=12, bold=False, space_after=6, caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for i, line in enumerate(text.split("\n")):
        r = p.add_run(("\n" if i else "") + (line.upper() if caps else line))
        r.font.size = Pt(size)
        r.bold = bold
        r.font.name = BODY_FONT
    return p


def _page_number_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def cover(doc: Document) -> None:
    _centred(doc, UNIVERSITY, 14, True, 2)
    _centred(doc, UNIVERSITY_ADDR, 12, False, 18)
    _centred(doc, "PROJECT PHASE-1 REPORT", 15, True, 18)
    _centred(doc, TITLE, 15, True, 4)
    _centred(doc, f'"{SUBTITLE}"', 13, True, 18)
    _centred(doc,
             "Submitted in partial fulfillment of the requirements for the Sixth Semester\n"
             "BACHELOR OF ENGINEERING\nin\nCOMPUTER SCIENCE AND ENGINEERING\n"
             "(IoT, Cyber-Security and Blockchain Technology)\n"
             f"for the Academic Year {YEAR}", 12, False, 18)
    _centred(doc, "Submitted by", 12, True, 8)
    t = doc.add_table(rows=len(TEAM), cols=2)
    t.alignment = 1
    for row, (name, usn) in zip(t.rows, TEAM):
        for cell, val, bold in ((row.cells[0], name, True), (row.cells[1], usn, False)):
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(2)
            r = cp.add_run(val); r.bold = bold; r.font.size = Pt(12); r.font.name = BODY_FONT
    doc.add_paragraph()
    _centred(doc, "Under the guidance of", 12, False, 4)
    _centred(doc, GUIDE, 13, True, 2)
    _centred(doc, GUIDE_TITLE, 12, False, 18)
    _centred(doc, DEPT, 12, True, 6)
    _centred(doc, COLLEGE, 13, True, 2)
    _centred(doc, COLLEGE_ADDR, 11, False, 0)
    doc.add_page_break()


def _rule(doc, width_hint: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "6")
    bt.set(qn("w:space"), "1"); bt.set(qn("w:color"), "000000")
    bd.append(bt); pr.append(bd)


def _para(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_after=8, spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = spacing
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.name = BODY_FONT
    return p


def _signature_row(doc, cols: list[tuple[str, str]]) -> None:
    """Signature blocks laid out in a borderless table so they align on the page."""
    t = doc.add_table(rows=2, cols=len(cols))
    for i, (role, name) in enumerate(cols):
        for row, text, bold in ((0, role, True), (1, name, False)):
            cp = t.rows[row].cells[i].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(2)
            r = cp.add_run(text); r.bold = bold
            r.font.size = Pt(11.5); r.font.name = BODY_FONT


def certificate(doc: Document) -> None:
    _centred(doc, COLLEGE, 13, True, 2)
    _centred(doc, "Bengaluru - 562157", 11, False, 4)
    _centred(doc, f"Department of {DEPT_SHORT}", 11.5, True, 24)
    _centred(doc, "CERTIFICATE", 15, True, 6)
    _rule(doc)
    names = ", ".join(f"{n} ({u})" for n, u in TEAM)
    _para(doc,
          f'This is to certify that the project work entitled "{TITLE}" carried out jointly by '
          f'{names}, are bonafide students of {COLLEGE_TC} in partial fulfillment for the '
          f'award of Bachelor of Engineering in the Department of {DEPT_SHORT} of the '
          f'Visvesvaraya Technological University, Belagavi during the year {YEAR_SHORT}. It is '
          f'certified that all corrections and suggestions indicated for internal assessment '
          f'have been incorporated in the report. The project report has been approved as it '
          f'satisfies the academic requirements in respect of the project work prescribed for '
          f'the said degree.', spacing=1.35)
    doc.add_paragraph()
    _signature_row(doc, [("Guide", GUIDE), ("Head of the Department", GUIDE),
                         ("Principal", PRINCIPAL)])
    _para(doc, "Submitted for the semester end examination (viva-voce) held on "
               "_______________________", space_after=14, spacing=1.2)
    _centred(doc, "External Viva", 12, True, 8)
    _signature_row(doc, [("Name of the Examiners", ""), ("Signature with Date", "")])
    for n in (1, 2):
        _para(doc, f"{n}. ______________________________"
                   f"                    ______________________________",
              size=11.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12, spacing=1.0)
    doc.add_page_break()


def declaration(doc: Document) -> None:
    _centred(doc, "DECLARATION BY THE CANDIDATES", 15, True, 6)
    _rule(doc)
    _para(doc,
          f'We, the undersigned, solemnly declare that the project work entitled "{TITLE}" is '
          f'based on our own research work carried out during the course of our study under the '
          f'supervision of {GUIDE}. We assert that the statements made and the conclusions drawn '
          f'are an outcome of our project work. We further certify that:')
    for item in (
        "The work contained in the report is original and has been done by us under the "
        "general supervision of our supervisor.",
        "The work has not been submitted to any other institution for any other "
        "degree, diploma or certificate in this university or any other university of "
        "India or abroad.",
        "We have followed the guidelines provided by the university in writing the report.",
        "Whenever we have used material - data, theoretical analysis or text - from other "
        "sources, we have given due credit to them in the text of the report and their "
        "details are provided in the references.",
    ):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(item); r.font.size = Pt(12); r.font.name = BODY_FONT
    doc.add_paragraph()
    _para(doc, "Date:", space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(doc, "Place: Bengaluru", space_after=18, align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(doc, "Project Associates:", bold=True, space_after=10,
          align=WD_ALIGN_PARAGRAPH.LEFT)
    t = doc.add_table(rows=len(TEAM), cols=3)
    for row, (name, usn) in zip(t.rows, TEAM):
        for cell, val, bold in ((row.cells[0], name, True), (row.cells[1], usn, False),
                                (row.cells[2], "________________", False)):
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(8)
            r = cp.add_run(val); r.bold = bold
            r.font.size = Pt(12); r.font.name = BODY_FONT
    doc.add_page_break()


def acknowledgement(doc: Document) -> None:
    _centred(doc, "ACKNOWLEDGEMENT", 15, True, 6)
    _rule(doc)
    for para in (
        "We would first like to thank our parents for their kind help, encouragement and "
        "moral support throughout the course of this work.",
        f"We thank {PRINCIPAL}, Principal, {COLLEGE_TC}, Bengaluru, for providing the "
        f"necessary support and facilities.",
        f"We would like to place on record our regards to {GUIDE}, {GUIDE_TITLE}, Department "
        f"of {DEPT_SHORT}, for her continued support.",
        f"We would like to thank our project coordinator {COORDINATOR}, {COORDINATOR_TITLE}, "
        f"for his support and coordination.",
        f"We would like to thank our project guide {GUIDE}, {GUIDE_TITLE}, for her continuous "
        f"support, valuable guidance and supervision towards the successful completion of "
        f"this project work.",
        "Finally, we thank all the members of the teaching and non-teaching staff of the "
        "department who assisted us during this work.",
    ):
        _para(doc, para)
    doc.add_paragraph()
    for name, usn in TEAM:
        _para(doc, f"{name}    {usn}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT,
              space_after=4, spacing=1.15)
    doc.add_page_break()


def abstract(doc: Document) -> None:
    _centred(doc, "ABSTRACT", 15, True, 6)
    _rule(doc)
    _para(doc, ABSTRACT)
    doc.add_page_break()


def abbreviations(doc: Document) -> None:
    _centred(doc, "LIST OF ABBREVIATIONS", 15, True, 6)
    _rule(doc)
    t = doc.add_table(rows=len(ABBREVIATIONS), cols=2)
    t.style = "Table Grid"
    for row, (abbr, meaning) in zip(t.rows, ABBREVIATIONS):
        for cell, val, bold in ((row.cells[0], abbr, True), (row.cells[1], meaning, False)):
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.line_spacing = 1.0
            r = cp.add_run(val); r.bold = bold
            r.font.size = Pt(11); r.font.name = BODY_FONT
    doc.add_page_break()


def _scan_captions(files: list[Path]) -> tuple[list[str], list[str]]:
    """Collect figure and table captions from the chapter sources.

    Generated up front from the Markdown rather than maintained by hand, so the
    lists cannot drift out of step with the chapters the way the template's
    stale List of Figures had."""
    figs, tabs = [], []
    fig_re = re.compile(r"^\*\*(Figure\s+[\d.]+)\*\*\s*[-\u2014]+\s*`[^`]+`\s*[-\u2014]+\s*(.+)$")
    tab_re = re.compile(r"^\*\*(Table\s+[\d.]+)\s*[-\u2014]+\s*(.+?)\*\*\s*$")
    tab_h = re.compile(r"^#{2,4}\s+(Table\s+[\d.]+)\s*[-\u2014]+\s*(.+)$")
    for f in files:
        if not f.exists():
            continue
        for line in f.read_text(encoding="utf-8").split("\n"):
            line = line.strip()
            m = fig_re.match(line)
            if m:
                figs.append(f"{m.group(1)}: {m.group(2).rstrip('.')}")
                continue
            m = tab_re.match(line) or tab_h.match(line)
            if m:
                tabs.append(f"{m.group(1)}: {m.group(2).rstrip('.').strip('*')}")
    # de-duplicate, keeping first occurrence order
    seen = set()
    figs = [x for x in figs if not (x in seen or seen.add(x))]
    seen = set()
    tabs = [x for x in tabs if not (x in seen or seen.add(x))]
    return figs, tabs


def _caption_list(doc: Document, title: str, items: list[str], empty: str) -> None:
    _centred(doc, title, 15, True, 6)
    _rule(doc)
    if not items:
        _para(doc, empty, size=11.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for item in items:
        label, _, rest = item.partition(":")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.1)
        r = p.add_run(label + "  "); r.bold = True
        r.font.size = Pt(11.5); r.font.name = BODY_FONT
        r = p.add_run(rest.strip())
        r.font.size = Pt(11.5); r.font.name = BODY_FONT
    doc.add_page_break()


def toc(doc: Document) -> None:
    """A real Word TOC field - Word fills it on F9 / print preview."""
    doc.add_heading("TABLE OF CONTENTS", level=1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = "Right-click and choose 'Update Field' to build the table of contents."
    run.append(t); fld.append(run); p._p.append(fld)
    doc.add_page_break()


# ------------------------------------------------------------- md -> docx

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[0-9a-z]+\])")


def _runs(par, text: str) -> None:
    """Bold, italic, inline code and citation markers, applied per run."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            par.add_run(piece[1:-1]).italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1]); r.font.name = MONO_FONT; r.font.size = Pt(10.5)
        else:
            par.add_run(piece)


def _table(doc, rows: list[str]) -> None:
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j, val in enumerate(row[: len(cells[0])]):
            p = t.rows[i].cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            _runs(p, val)
            for r in p.runs:
                r.font.size = Pt(10.5); r.font.name = BODY_FONT
                if i == 0:
                    r.bold = True
    doc.add_paragraph()


def render(doc: Document, md: str, fig_counter: list[int]) -> None:
    lines = md.split("\n")
    i, buf = 0, []

    def flush():
        nonlocal buf
        if buf:
            text = " ".join(x.strip() for x in buf).strip()
            if text:
                _runs(doc.add_paragraph(), text)
            buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run("\n".join(code)); r.font.name = MONO_FONT; r.font.size = Pt(9.5)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i]); i += 1
            _table(doc, rows)
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            flush()
            level, text = len(m.group(1)), m.group(2).strip()
            if level == 1 and re.match(r"^Chapter\s+\d", text, re.I):
                doc.add_page_break()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Two figure conventions are accepted: standard Markdown images, and the
        # chapters' own "**Figure N.N** - `path` - caption" caption lines.
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        fig_line = None
        if not m:
            fig_line = re.match(
                r"^\*\*(Figure\s+[\d.]+)\*\*\s*[-\u2014]+\s*`([^`]*docs/figures/[^`]+)`\s*[-\u2014]+\s*(.*)$",
                stripped)
        if fig_line:
            flush()
            label, src, cap = fig_line.group(1), fig_line.group(2), fig_line.group(3)
            # a caption may wrap onto following lines until a blank line
            j = i + 1
            while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(("#", "|", "**Figure")):
                cap += " " + lines[j].strip(); j += 1
            path = REPO / src
            if path.suffix == ".svg":
                path = path.with_suffix(".png")
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.1))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_counter[0] += 1
                c = doc.add_paragraph()
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraph_format.space_after = Pt(12)
                r = c.add_run(f"{label}: {cap.rstrip('.')}")
                r.font.size = Pt(10.5); r.italic = True
            else:
                _runs(doc.add_paragraph(), f"[missing figure: {src}]")
            i = j
            continue

        if m:
            flush()
            caption_text, src = m.group(1), m.group(2)
            path = (REPO / src) if not src.startswith("/") else Path(src)
            if path.suffix == ".svg":
                path = path.with_suffix(".png")
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_counter[0] += 1
                c = doc.add_paragraph()
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = c.add_run(f"Fig. {fig_counter[0]}: {caption_text}")
                r.font.size = Pt(10.5); r.italic = True
            else:
                _runs(doc.add_paragraph(), f"[missing figure: {src}]")
            i += 1
            continue

        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            flush()
            indent, marker, text = m.group(1), m.group(2), m.group(3)
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.35 + 0.3 * (len(indent) // 2))
            p.paragraph_format.space_after = Pt(3)
            _runs(p, text)
            i += 1
            continue

        if not stripped or set(stripped) <= set("-*_") and len(stripped) >= 3:
            flush()
            i += 1
            continue

        buf.append(line)
        i += 1

    flush()


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--out", type=Path,
                    default=REPORT_DIR / "EdgeGrid_Phase1_Report.docx")
    ap.add_argument("--chapters", nargs="*", default=None)
    args = ap.parse_args(argv)

    doc = Document()
    for s in doc.sections:
        # A4 with a wider left margin for binding - the VTU convention.
        s.page_width, s.page_height = Mm(210), Mm(297)
        s.top_margin = s.bottom_margin = Mm(25)
        s.left_margin = Mm(32); s.right_margin = Mm(25)
    _style(doc)
    _page_number_footer(doc.sections[0])

    files = args.chapters or CHAPTER_FILES
    chapter_paths = [REPORT_DIR / f for f in files]
    fig_captions, tab_captions = _scan_captions(chapter_paths)

    # front matter, in the order the department's template prescribes
    cover(doc)
    certificate(doc)
    declaration(doc)
    acknowledgement(doc)
    toc(doc)
    abstract(doc)
    abbreviations(doc)
    _caption_list(doc, "LIST OF FIGURES", fig_captions,
                  "No figures are referenced in the chapters included in this build.")
    _caption_list(doc, "LIST OF TABLES", tab_captions,
                  "No tables are referenced in the chapters included in this build.")

    fig_counter = [0]
    included, missing = [], []
    for name in files:
        path = REPORT_DIR / name
        if not path.exists():
            missing.append(name)
            continue
        render(doc, path.read_text(encoding="utf-8"), fig_counter)
        included.append(name)

    refs = HERE / "REFERENCES.md"
    if refs.exists():
        doc.add_page_break()
        render(doc, refs.read_text(encoding="utf-8"), fig_counter)
        included.append("REFERENCES.md")

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)

    words = sum(len((REPORT_DIR / f).read_text().split()) for f in included
                if (REPORT_DIR / f).exists())
    print(f"wrote {args.out}")
    print(f"  chapters : {', '.join(included)}")
    print(f"  front    : cover, certificate, declaration, acknowledgement, contents, "
          f"abstract, {len(ABBREVIATIONS)} abbreviations, "
          f"{len(fig_captions)} figures listed, {len(tab_captions)} tables listed")
    print(f"  figures  : {fig_counter[0]} embedded")
    print(f"  ~words   : {words:,}")
    if missing:
        print(f"  PENDING  : {', '.join(missing)} (not yet written - run the experiments first)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
